"""
Word 报告输出引擎
=================
将计算参数和图表填入 Word 模板 (.docx)。

定位方式（优先级）:
  1. 内容控件 (Content Control / SDT) → 按 tag 匹配
  2. 书签 (Bookmark) → 按名定位插图片
  3. 表格 (Table) → 列头识别（复用 classify_header）
  4. 占位符 ({{variable}}) → 搜索替换

所有定位信息在 fill 时收集并缓存，可在 GUI 预览。
"""
from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ═══════════════════════════════════════════════════════════════
# 数据类
# ═══════════════════════════════════════════════════════════════

class WordTemplateInfo:
    """模板扫描结果 — 用于 GUI 预览。"""
    def __init__(self):
        self.tables: List[dict] = []        # [{sheet_index, header_row, columns: [{col_type, text}]}]
        self.bookmarks: List[str] = []       # [bookmark_name, ...]
        self.content_controls: List[str] = []  # [tag_name, ...]
        self.placeholders: List[str] = []    # [var_name, ...]

    @property
    def has_content(self) -> bool:
        return bool(self.tables or self.bookmarks or self.content_controls or self.placeholders)


# ═══════════════════════════════════════════════════════════════
# Word 报告填充器
# ═══════════════════════════════════════════════════════════════

class WordReporter:
    """将天线参数报告填入 Word 模板。"""

    def __init__(self, template_path: str):
        self._doc = Document(template_path)
        self._template_path = template_path
        self._info = WordTemplateInfo()
        self._filled_count = 0  # 统计填入项

    # ── 模板扫描 ──────────────────────────────────────────

    def scan(self) -> WordTemplateInfo:
        """扫描模板，收集所有可填入位置的信息。"""
        self._scan_tables()
        self._scan_bookmarks()
        self._scan_content_controls()
        self._scan_placeholders()
        return self._info

    def _scan_tables(self):
        """扫描所有表格的列头。"""
        for ti, table in enumerate(self._doc.tables):
            if not table.rows:
                continue
            header_cells = table.rows[0].cells
            columns = []
            for ci, cell in enumerate(header_cells):
                text = cell.text.strip()
                if text:
                    from src.column_mapping import classify_header
                    col_type = classify_header(text)
                    columns.append({"index": ci, "text": text, "col_type": col_type})
                else:
                    columns.append({"index": ci, "text": "", "col_type": "unknown"})
            self._info.tables.append({
                "sheet_index": ti,
                "header_row": 0,
                "columns": columns,
            })

    def _scan_bookmarks(self):
        """扫描所有书签。"""
        for bookmark in self._doc.element.body.iter(qn('w:bookmarkStart')):
            name = bookmark.get(qn('w:name'), '')
            if name and name not in self._info.bookmarks:
                self._info.bookmarks.append(name)

    def _scan_content_controls(self):
        """扫描所有内容控件 (Structured Document Tags) 的 tag。"""
        for sdt in self._doc.element.body.iter(qn('w:sdt')):
            tag_el = sdt.find(qn('w:sdtPr') + '/' + qn('w:tag'))
            if tag_el is not None:
                tag_val = tag_el.get(qn('w:val'), '')
                if tag_val:
                    self._info.content_controls.append(tag_val)

    def _scan_placeholders(self):
        """扫描 {{variable}} 占位符。"""
        text = self._doc_text()
        for m in re.finditer(r'\{\{(\w+)\}\}', text):
            var = m.group(1)
            if var not in self._info.placeholders:
                self._info.placeholders.append(var)

    def _doc_text(self) -> str:
        """获取文档全部文字（含表格）。"""
        texts = []
        for para in self._doc.paragraphs:
            texts.append(para.text)
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return "\n".join(texts)

    # ── 填充表格 ──────────────────────────────────────────

    def fill_tables(self, sheet_results: Dict[str, List[Dict[str, Any]]],
                    progress_callback=None) -> int:
        """将计算结果填入模板中的表格。

        按顺序匹配:
          1. 表格列头检测 col_type
          2. sheet_name 匹配工作表名

        Returns:
            填入的行数
        """
        if not self._info.tables:
            self._scan_tables()

        total_rows = 0

        for ti, tinfo in enumerate(self._info.tables):
            if ti >= len(self._doc.tables):
                break
            table = self._doc.tables[ti]

            # 构建列映射: col_type → column index
            col_map: Dict[str, int] = {}
            for col in tinfo["columns"]:
                ct = col["col_type"]
                if ct != "unknown" and ct not in col_map:
                    col_map[ct] = col["index"]

            if not col_map:
                continue  # 没有可识别的列

            # 找匹配的数据
            sheet_names = list(sheet_results.keys())
            rows_data = None
            if ti < len(sheet_names):
                rows_data = sheet_results.get(sheet_names[ti])

            if not rows_data:
                continue

            header_row_idx = tinfo["header_row"]
            for ri, row_dict in enumerate(rows_data):
                row_idx = header_row_idx + 1 + ri
                # 确保表格有足够行
                while len(table.rows) <= row_idx:
                    table.add_row()

                for col_type, col_idx in col_map.items():
                    value = row_dict.get(col_type)
                    if value is None:
                        continue
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(round(value, 6) if isinstance(value, float) else value)

                total_rows += 1
                if progress_callback:
                    progress_callback(ri + 1, len(rows_data), f"填表 {ti+1}")

        self._filled_count += total_rows
        return total_rows

    # ── 填充内容控件 ──────────────────────────────────────

    def fill_content_controls(self, data: Dict[str, Any]) -> int:
        """按 tag 名匹配内容控件并填入值。

        data: {tag_name: value, ...}
        """
        count = 0
        for sdt in self._doc.element.body.iter(qn('w:sdt')):
            tag_el = sdt.find(qn('w:sdtPr') + '/' + qn('w:tag'))
            if tag_el is None:
                continue
            tag_val = tag_el.get(qn('w:val'), '')
            if tag_val not in data:
                continue

            value = data[tag_val]
            text = str(round(value, 6) if isinstance(value, float) else value)

            # 找到 sdtContent 中的第一个 run
            content = sdt.find(qn('w:sdtContent'))
            if content is None:
                continue
            run = content.find(qn('w:r'))
            if run is not None:
                t = run.find(qn('w:t'))
                if t is not None:
                    t.text = text
                    count += 1
                else:
                    # 没有 t → 创建一个
                    new_t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{text}</w:t>')
                    run.append(new_t)
                    count += 1

        self._filled_count += count
        return count

    # ── 填充占位符 ────────────────────────────────────────

    def fill_placeholders(self, data: Dict[str, Any]) -> int:
        """搜索替换 {{variable}} 占位符。

        data: {var_name: value, ...}
        """
        count = 0
        for para in self._doc.paragraphs:
            for run in para.runs:
                for key, value in data.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in run.text:
                        val = str(round(value, 6) if isinstance(value, float) else value)
                        run.text = run.text.replace(placeholder, val)
                        count += 1

        # 也替换表格中的占位符
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, value in data.items():
                                placeholder = '{{' + key + '}}'
                                if placeholder in run.text:
                                    val = str(round(value, 6) if isinstance(value, float) else value)
                                    run.text = run.text.replace(placeholder, val)
                                    count += 1

        self._filled_count += count
        return count

    # ── 插入图片（书签定位） ────────────────────────────────

    def insert_images_at_bookmarks(self, bookmark_images: Dict[str, bytes],
                                    max_width_inches: float = 5.5) -> int:
        """在书签位置插入图片（使用 tempfile + add_picture）。"""
        import tempfile
        from docx.shared import Inches

        body = self._doc.element.body
        bookmark_paras: Dict[str, int] = {}
        for el in body.iter():
            if el.tag == qn('w:bookmarkStart'):
                name = el.get(qn('w:name'), '')
                if name in bookmark_images:
                    parent = el.getparent()
                    if parent is not None:
                        for pi, para in enumerate(self._doc.paragraphs):
                            if para._element is parent:
                                bookmark_paras[name] = pi
                                break

        count = 0
        for name, para_idx in bookmark_paras.items():
            png_bytes = bookmark_images[name]
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp.write(png_bytes)
            tmp.close()
            try:
                para = self._doc.paragraphs[para_idx]
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                img_para = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
                para._element.addnext(img_para)

                handle_doc = Document()
                new_run = handle_doc.add_paragraph().add_run()
                new_run.add_picture(tmp.name, width=Inches(max_width_inches))

                for child in new_run._r:
                    img_para.append(child)
                count += 1
            finally:
                try:
                    os.unlink(tmp.name)
                except OSError:
                    pass

        self._filled_count += count
        return count

    # ── 循环域 + 元数据 ──────────────────────────────────

    _RE_LOOP_START = re.compile(r'\{\{loop_start_(\w+)\}\}')
    _RE_LOOP_END = re.compile(r'\{\{loop_end_(\w+)\}\}')

    def _expand_loops(self) -> Dict[str, tuple]:
        """扫描段落和表格中 {{loop_start_<key>}} / {{loop_end_<key>}} 配对。"""
        self._loop_groups = {}
        open_entry: Dict[str, int] = {}

        def _scan_paragraphs(paragraphs):
            for i, para in enumerate(paragraphs):
                text = para.text.strip()
                m_start = self._RE_LOOP_START.match(text)
                m_end = self._RE_LOOP_END.match(text)
                if m_start:
                    open_entry[m_start.group(1)] = i
                elif m_end:
                    key = m_end.group(1)
                    if key in open_entry:
                        self._loop_groups[key] = (open_entry.pop(key), i)

        # 扫描正文段落
        _scan_paragraphs(self._doc.paragraphs)
        # 扫描所有表格单元格
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _scan_paragraphs(cell.paragraphs)
        return dict(self._loop_groups)

    def fill_metadata(self, metadata: Dict[str, Any]) -> int:
        """填充元数据: 内容控件 + 占位符，统一入口。"""
        return self.fill_content_controls(metadata) + self.fill_placeholders(metadata)

    # ── 通用填充 ──────────────────────────────────────────

    def fill_all(self, sheet_results: Dict[str, List[Dict]],
                 single_values: Optional[Dict[str, Any]] = None,
                 bookmark_images: Optional[Dict[str, bytes]] = None,
                 progress_callback=None) -> dict:
        """执行全部填充操作。

        Returns:
            {"tables": N, "controls": N, "placeholders": N, "images": N}
        """
        result = {"tables": 0, "controls": 0, "placeholders": 0, "images": 0}

        result["tables"] = self.fill_tables(sheet_results, progress_callback)

        if single_values:
            result["controls"] = self.fill_content_controls(single_values)
            result["placeholders"] = self.fill_placeholders(single_values)

        if bookmark_images:
            result["images"] = self.insert_images_at_bookmarks(bookmark_images)

        return result

    # ── 输出 ──────────────────────────────────────────────

    def save(self, output_path: str) -> str:
        """保存填充后的文档。"""
        os.makedirs(str(Path(output_path).parent), exist_ok=True)
        self._doc.save(output_path)
        return output_path

    @property
    def template_info(self) -> WordTemplateInfo:
        return self._info

    @property
    def filled_count(self) -> int:
        return self._filled_count
