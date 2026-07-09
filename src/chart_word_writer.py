"""
图表 Word 报告输出
==================
将任意图表 PNG 图片写入 Word 文档 (.docx)。

支持多组图片（如 Gain Azimuth / AR Azimuth / 3D Gain 等），
每组可以有自己的标题前缀。A4 竖版，每行 N 列（默认 2），
图片宽度可配置为列宽的百分比。

接口:
  write_chart_word_report(image_groups, output_path, ...)

image_groups 结构:
  {
      "Gain Azimuth Cut": {1154.0: BytesIO, 1155.0: BytesIO, ...},
      "AR Azimuth Cut":   {1176.0: BytesIO, ...},
  }
  每个组按频点排序输出，组间顺序由传入 dict 顺序决定。
"""

from __future__ import annotations

import io
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def write_chart_word_report(
    image_groups: dict[str, dict[float, io.BytesIO]],
    output_path: str,
    antenna_name: str = "",
    angles_str: str = "",
    layout_columns: int = 2,
    image_width_pct: int = 90,
    label_order: list[str] | None = None,
    layout_mode: str = "side_by_side",
    show_heading: bool = False,
    show_caption: bool = False,
) -> None:
    """将多组图表图片写入 Word 文档。

    Args:
        image_groups: {组名: {频率MHz: PNG BytesIO}}.
        output_path: 输出 .docx 路径。
        antenna_name: 天线名，用于标题。
        angles_str: 角度描述串（如 "60°, 70°, 80°, 90°"），用于题注。
        layout_columns: 每行列数 (1 或 2)。
        image_width_pct: 图片宽度占列宽的百分比 (10-100)。
        label_order: 图表组输出顺序列表（如按 ChartInstance.sort_order）。
        layout_mode: "side_by_side", "by_freq", "by_type"。
        show_caption: 是否显示图片上方题注。
    """
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    content_width_cm = 17.0  # A4 21cm - 2×2cm margins
    col_width_cm = content_width_cm / max(1, layout_columns)
    img_width = Cm(col_width_cm * image_width_pct / 100.0)

    # 按 label_order 排序输出（若提供）
    if label_order:
        ordered_groups = [(name, image_groups[name]) for name in label_order
                          if name in image_groups and image_groups[name]]
        # 追加未在 order 中的 group
        seen = set(label_order)
        for name, images in image_groups.items():
            if name not in seen and images:
                ordered_groups.append((name, images))
    else:
        ordered_groups = [(name, images) for name, images in image_groups.items() if images]

    if layout_mode == "by_freq":
        # by_freq: 每频点一行, 展示所有图表类型
        _write_by_freq(doc, ordered_groups, antenna_name, angles_str, img_width, layout_columns, show_heading, show_caption)
    elif layout_mode == "by_type":
        # by_type: 每图表类型一节, 展示所有频点 (等同于默认逐组排列)
        _write_by_type(doc, ordered_groups, antenna_name, angles_str, img_width, layout_columns, show_heading, show_caption)
    else:
        # side_by_side / default: 逐组排列
        _write_by_type(doc, ordered_groups, antenna_name, angles_str, img_width, layout_columns, show_heading, show_caption)

    # 避免覆盖: 同名文件自动加 _1, _2, ...
    base, ext = os.path.splitext(output_path)
    n = 1
    while os.path.exists(output_path):
        output_path = f"{base}_{n}{ext}"
        n += 1

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)


def write_chart_word_report_by_freq(
    freq_pairs: dict[float, dict[str, io.BytesIO]],
    pair_order: list[str],
    pair_labels: dict[str, str],
    output_path: str,
    antenna_name: str = "",
    image_width_pct: int = 90,
    extra_groups: dict[str, dict[float, io.BytesIO]] = None,
    extra_angles: str = "",
    show_caption: bool = True,
) -> None:
    """按频点排列图表: 每频点一行 N 张图并排。"""
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    content_width_cm = 17.0  # A4 21cm - 2×2cm margins
    n_cols = max(1, len(pair_order))
    col_width_cm = content_width_cm / n_cols
    img_width = Cm(col_width_cm * image_width_pct / 100.0)
    freqs = sorted(freq_pairs.keys())

    for freq in freqs:
        images = freq_pairs[freq]
        if n_cols == 1:
            key = pair_order[0]
            buf = images.get(key)
            if buf:
                _add_single_image(doc, buf,
                    f"{freq:.0f} MHz — {pair_labels.get(key, key)}" if show_caption else "",
                    width=img_width, show_caption=show_caption)
        else:
            table = doc.add_table(rows=1, cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, key in enumerate(pair_order):
                buf = images.get(key)
                if buf:
                    _add_cell_image(table.cell(0, ci), buf,
                        f"{freq:.0f} MHz — {pair_labels.get(key, key)}" if show_caption else "",
                        width=img_width, show_caption=show_caption)

    if extra_groups:
        for group_name, images in extra_groups.items():
            if not images: continue
            if show_heading:
                heading = doc.add_heading(group_name, level=1); heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            freqs_extra = sorted(images.keys())
            _write_image_grid(doc, images, freqs_extra, antenna_name,
                              group_name, extra_angles, img_width, 1,
                              show_caption=show_caption)

    # 避免覆盖: 同名文件自动加 _1, _2, ...
    base, ext = os.path.splitext(output_path)
    n = 1
    while os.path.exists(output_path):
        output_path = f"{base}_{n}{ext}"
        n += 1

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)


def _make_caption(antenna_name: str, freq_mhz: float, group_name: str,
                  angles_str: str) -> str:
    """生成图片题注。"""
    parts = []
    if antenna_name:
        parts.append(antenna_name)
    parts.append(f"{freq_mhz:.0f} MHz")
    parts.append(group_name)
    if angles_str:
        parts.append(f"(θ={angles_str})")
    return " — ".join(parts)


def _add_single_image(doc: Document, img_buf: io.BytesIO, caption: str,
                      width: object = Cm(7.5), show_caption: bool = True) -> None:
    """添加一张带上方题注的居中图片。"""
    if show_caption and caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.bold = True
        run.font.size = Pt(9)
        cap_para.paragraph_format.space_after = Pt(2)
    img_buf.seek(0)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = img_para.add_run()
    run_img.add_picture(img_buf, width=width)
    img_para.paragraph_format.space_after = Pt(4)


def _add_cell_image(cell, img_buf: io.BytesIO, caption: str,
                    width: object = Cm(7.5), show_caption: bool = True) -> None:
    """在表格单元格中添加图片 + 上方题注。"""
    if show_caption and caption:
        cap_para = cell.paragraphs[0]
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.bold = True
        run.font.size = Pt(7)
        cap_para.paragraph_format.space_after = Pt(2)
    img_buf.seek(0)
    img_para = cell.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = img_para.add_run()
    run_img.add_picture(img_buf, width=width)


def _write_image_grid(
    doc: Document,
    images: dict[float, io.BytesIO],
    freqs: list[float],
    antenna_name: str,
    group_name: str,
    angles_str: str,
    img_width: object,
    columns: int = 2,
    show_caption: bool = True,
) -> None:
    """按频点排列图片，每行 N 列。"""
    def _bufs(freq):
        v = images.get(freq)
        if v is None: return []
        return v if isinstance(v, list) else [v]

    if columns == 1:
        for freq in freqs:
            for buf in _bufs(freq):
                cap = _make_caption(antenna_name, freq, group_name, angles_str) if show_caption else ""
                _add_single_image(doc, buf, cap, width=img_width, show_caption=show_caption)
    else:
        for i in range(0, len(freqs), columns):
            row_freqs = freqs[i:i + columns]
            if len(row_freqs) == 1:
                for buf in _bufs(row_freqs[0]):
                    cap = _make_caption(antenna_name, row_freqs[0], group_name, angles_str) if show_caption else ""
                    _add_single_image(doc, buf, cap, width=img_width, show_caption=show_caption)
            else:
                table = doc.add_table(rows=1, cols=len(row_freqs))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, freq in enumerate(row_freqs):
                    for buf in _bufs(freq):
                        cap = _make_caption(antenna_name, freq, group_name, angles_str) if show_caption else ""
                        _add_cell_image(table.cell(0, j), buf, cap, width=img_width, show_caption=show_caption)
                # 图片间不加空段 — 紧凑排列


def _write_by_freq(doc, ordered_groups, antenna_name, angles_str, img_width, columns, show_heading, show_caption):
    """按频点布局: 每个频点展示该频点下的所有图表类型。"""
    # 收集所有频点 + 所有图表类型
    all_freqs = set()
    for _, images in ordered_groups:
        all_freqs.update(images.keys())
    all_freqs = sorted(all_freqs)

    group_names = [name for name, _ in ordered_groups]

    for freq in all_freqs:
        # 收集该频点的所有图片
        freq_images = {}
        for group_name, images in ordered_groups:
            if freq in images:
                freq_images[group_name] = {freq: images[freq]}

        if not freq_images:
            continue

        # 用表格排列 (每行最多 columns 列)
        entries = list(freq_images.items())
        for i in range(0, len(entries), columns):
            row_entries = entries[i:i + columns]
            if len(row_entries) == 1:
                gn, imgs = row_entries[0]
                cap = _make_caption(antenna_name, freq, gn, angles_str) if show_caption else ""
                _add_single_image(doc, imgs[freq], cap, width=img_width, show_caption=show_caption)
            else:
                table = doc.add_table(rows=1, cols=len(row_entries))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, (gn, imgs) in enumerate(row_entries):
                    cap = _make_caption(antenna_name, freq, gn, angles_str) if show_caption else ""
                    _add_cell_image(table.cell(0, j), imgs[freq], cap, width=img_width, show_caption=show_caption)


def _write_by_type(doc, ordered_groups, antenna_name, angles_str, img_width, columns, show_heading, show_caption):
    """按图表类型布局: 每种图表类型展示其所有频点的图片。"""
    for group_name, images in ordered_groups:
        if show_heading:
            heading = doc.add_heading(group_name, level=1); heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        freqs = sorted(images.keys())
        _write_image_grid(doc, images, freqs, antenna_name, group_name,
                          angles_str, img_width, columns, show_caption=show_caption)
